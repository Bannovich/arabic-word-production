from __future__ import annotations

import importlib.util
from pathlib import Path
from tempfile import TemporaryDirectory
import unittest

from docx import Document
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[1]
AUDIT_SCRIPT = ROOT / "scripts" / "audit_docx.py"


def load_auditor(testcase: unittest.TestCase):
    if not AUDIT_SCRIPT.exists():
        testcase.fail("auditor_missing: scripts/audit_docx.py does not exist")
    spec = importlib.util.spec_from_file_location("arabic_word_audit_docx", AUDIT_SCRIPT)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module.audit_docx


class AuditorTests(unittest.TestCase):
    def test_reports_stable_ids_for_missing_rtl_table_direction_and_placeholder(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "malformed.docx"
            document = Document()
            document.add_paragraph("هذه فقرة عربية بلا اتجاه صريح.")
            table = document.add_table(rows=2, cols=7)
            for index, cell in enumerate(table.rows[0].cells):
                cell.text = f"عمود {index + 1}"
            for index, cell in enumerate(table.rows[1].cells):
                cell.text = f"قيمة {index + 1}"
            document.add_paragraph("TBD")
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-RTL-001", ids)
            self.assertIn("ERR-TABLE-001", ids)
            self.assertIn("ERR-CONTENT-001", ids)
            self.assertFalse(result["passed"])

    def test_reports_missing_source_content_when_model_is_supplied(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "incomplete.docx"
            document = Document()
            paragraph = document.add_paragraph("عنوان موجود")
            ppr = paragraph._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            ppr.append(bidi)
            document.save(path)
            model = {
                "title": "عنوان موجود",
                "blocks": [{"type": "paragraph", "text": "فقرة مفقودة من المستند"}],
            }

            try:
                result = audit_docx(path, model)
            except TypeError:
                self.fail("model_comparison_missing: audit_docx does not accept a source model")

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-CONTENT-003", ids)
            self.assertFalse(result["passed"])

    def test_accepts_source_content_separated_by_word_line_breaks(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "line-break.docx"
            document = Document()
            paragraph = document.add_paragraph()
            ppr = paragraph._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            ppr.append(bidi)
            run = paragraph.add_run("السطر الأول")
            run.add_break()
            run.add_text("السطر الثاني")
            document.save(path)

            model = {
                "title": "",
                "blocks": [
                    {"type": "paragraph", "text": "السطر الأول\nالسطر الثاني"}
                ],
            }
            result = audit_docx(path, model)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertNotIn("ERR-CONTENT-003", ids)

    def test_reports_missing_headers_footers_and_dynamic_page_fields(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "missing-page-chrome.docx"
            document = Document()
            paragraph = document.add_paragraph("فقرة مكتملة الاتجاه.")
            ppr = paragraph._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            ppr.append(bidi)
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-SECTION-001", ids)
            self.assertFalse(result["passed"])

    def test_reports_trailing_edge_alignment_on_an_rtl_paragraph(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "rtl-trailing-edge.docx"
            document = Document()
            paragraph = document.add_paragraph("عنوان عربي يجب أن يظهر يمينًا.")
            ppr = paragraph._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            bidi = OxmlElement("w:bidi")
            bidi.set(qn("w:val"), "1")
            ppr.append(bidi)
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "right")
            ppr.append(jc)
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-RTL-003", ids)

    def test_reports_an_rtl_mixed_page_field_footer(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "rtl-page-footer.docx"
            document = Document()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            body = document.add_paragraph("محتوى عربي.")
            body_bidi = OxmlElement("w:bidi")
            body_bidi.set(qn("w:val"), "1")
            body._p.get_or_add_pPr().append(body_bidi)

            document.sections[0].header.paragraphs[0].text = "Header"
            footer = document.sections[0].footer.paragraphs[0]
            footer.text = "صفحة "
            footer_bidi = OxmlElement("w:bidi")
            footer_bidi.set(qn("w:val"), "1")
            footer._p.get_or_add_pPr().append(footer_bidi)
            footer_jc = OxmlElement("w:jc")
            footer_jc.set(qn("w:val"), "center")
            footer._p.get_or_add_pPr().append(footer_jc)
            for field_name in ("PAGE", "NUMPAGES"):
                run = footer.add_run()
                instruction = OxmlElement("w:instrText")
                instruction.text = f" {field_name} "
                run._r.append(instruction)
            footer.add_run(" من ")
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-SECTION-001", ids)

    def test_reports_update_fields_on_open(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "update-fields-on-open.docx"
            document = Document()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            update_fields = OxmlElement("w:updateFields")
            update_fields.set(qn("w:val"), "1")
            document.settings.element.append(update_fields)
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertIn("ERR-FIELD-001", ids)

    def test_accepts_direction_inherited_from_paragraph_styles_after_word_roundtrip(self):
        audit_docx = load_auditor(self)
        with TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "word-style-inheritance.docx"
            document = Document()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            rtl_style = document.styles.add_style(
                "Roundtrip Arabic RTL",
                WD_STYLE_TYPE.PARAGRAPH,
            )
            rtl_bidi = OxmlElement("w:bidi")
            rtl_style.element.get_or_add_pPr().append(rtl_bidi)
            document.add_paragraph("فقرة عربية ترث اتجاهها من النمط.", rtl_style)

            document.sections[0].header.paragraphs[0].text = "Header"
            footer = document.sections[0].footer.paragraphs[0]
            footer.style = document.styles["Normal"]
            footer_jc = OxmlElement("w:jc")
            footer_jc.set(qn("w:val"), "center")
            footer._p.get_or_add_pPr().append(footer_jc)
            for field_name in ("PAGE", "NUMPAGES"):
                if footer.runs:
                    footer.add_run(" / ")
                run = footer.add_run()
                instruction = OxmlElement("w:instrText")
                instruction.text = f" {field_name} "
                run._r.append(instruction)
            document.save(path)

            result = audit_docx(path)

            ids = {finding["id"] for finding in result["findings"]}
            self.assertNotIn("ERR-RTL-001", ids)
            self.assertNotIn("ERR-SECTION-001", ids)


if __name__ == "__main__":
    unittest.main()
