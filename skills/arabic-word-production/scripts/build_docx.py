from __future__ import annotations

import argparse
import json
import re
import time
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ARABIC_RE = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")
URL_RE = re.compile(r"https?://\S+")


def _set_on_off(parent, tag: str, enabled: bool) -> None:
    element = parent.find(qn(tag))
    if element is None:
        element = OxmlElement(tag)
        parent.append(element)
    element.set(qn("w:val"), "1" if enabled else "0")


def _set_justification(parent, value: str) -> None:
    element = parent.find(qn("w:jc"))
    if element is None:
        element = OxmlElement("w:jc")
        parent.append(element)
    element.set(qn("w:val"), value)


def _set_paragraph_direction(paragraph, rtl: bool) -> None:
    properties = paragraph._p.get_or_add_pPr()
    _set_on_off(properties, "w:bidi", rtl)
    _set_justification(properties, "start")


def _set_run_direction(run, rtl: bool) -> None:
    _set_on_off(run._r.get_or_add_rPr(), "w:rtl", rtl)


def _add_directional_runs(paragraph, text: str, base_rtl: bool) -> None:
    cursor = 0
    for match in URL_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_run_direction(run, base_rtl)
        run = paragraph.add_run(match.group(0))
        _set_run_direction(run, False)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_direction(run, base_rtl)
    if not text:
        run = paragraph.add_run("")
        _set_run_direction(run, base_rtl)


def _is_rtl(text: str, explicit: str = "auto") -> bool:
    if explicit not in {"auto", "rtl", "ltr"}:
        raise ValueError(f"Unsupported direction: {explicit}")
    if explicit != "auto":
        return explicit == "rtl"
    return bool(ARABIC_RE.search(text))


def _add_text_paragraph(
    document: Document,
    text: str,
    direction: str = "auto",
    style_name: str | None = None,
):
    paragraph = document.add_paragraph()
    if style_name:
        paragraph.style = style_name
    rtl = _is_rtl(text, direction)
    _set_paragraph_direction(paragraph, rtl)
    _add_directional_runs(paragraph, text, rtl)
    return paragraph


def _set_style_font(style, name: str, size: float, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), name)


def _style(document: Document, name: str, base: str = "Normal"):
    styles = document.styles
    try:
        style = styles[name]
    except KeyError:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base and name != base:
        style.base_style = styles[base]
    return style


def _configure_style(
    document: Document,
    name: str,
    *,
    base: str = "Normal",
    rtl: bool,
    font: str,
    size: float,
    bold: bool = False,
    before: float = 0,
    after: float = 6,
    keep_with_next: bool = False,
    color: str | None = None,
):
    style = _style(document, name, base)
    _set_style_font(style, font, size, bold)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = keep_with_next
    style.paragraph_format.widow_control = True
    style_properties = style._element.get_or_add_pPr()
    _set_on_off(style_properties, "w:bidi", rtl)
    _set_justification(style_properties, "start")
    return style


def _configure_styles(document: Document) -> None:
    _configure_style(document, "Arabic Title RTL", rtl=True, font="Arial", size=26, bold=True, after=18, keep_with_next=True, color="17365D")
    _configure_style(document, "Arabic Heading 1 RTL", rtl=True, font="Arial", size=18, bold=True, before=14, after=7, keep_with_next=True, color="1F4E78")
    _configure_style(document, "Arabic Heading 2 RTL", rtl=True, font="Arial", size=14, bold=True, before=10, after=5, keep_with_next=True, color="2F5597")
    _configure_style(document, "Arabic Body RTL", rtl=True, font="Arial", size=11, after=7)
    _configure_style(document, "Arabic Mixed Body RTL", rtl=True, font="Arial", size=11, after=7)
    _configure_style(document, "English Body LTR", rtl=False, font="Aptos", size=10.5, after=7)
    _configure_style(document, "English Technical LTR", rtl=False, font="Consolas", size=9.5, after=7)
    _configure_style(document, "Source URL LTR", rtl=False, font="Aptos", size=9, after=4, color="0563C1")
    _configure_style(document, "Arabic Table Header RTL", rtl=True, font="Arial", size=9, bold=True, after=0)
    _configure_style(document, "Arabic Table Body RTL", rtl=True, font="Arial", size=8.5, after=0)
    _configure_style(document, "Arabic List Number RTL", base="List Number", rtl=True, font="Arial", size=11, after=4)
    _configure_style(document, "Arabic List Bullet RTL", base="List Bullet", rtl=True, font="Arial", size=11, after=4)
    _configure_style(document, "English List Number LTR", base="List Number", rtl=False, font="Aptos", size=10.5, after=4)
    _configure_style(document, "English List Bullet LTR", base="List Bullet", rtl=False, font="Aptos", size=10.5, after=4)
    _configure_style(document, "Note", rtl=True, font="Arial", size=10.5, after=8, color="385723")
    _configure_style(document, "Warning", rtl=True, font="Arial", size=10.5, bold=True, after=8, color="9C0006")
    _configure_style(document, "Caption", rtl=True, font="Arial", size=9, after=8, color="595959")


def _configure_modern_word_compatibility(document: Document) -> None:
    settings = document.settings._element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    compatibility_mode = None
    for node in compat.findall(qn("w:compatSetting")):
        if node.get(qn("w:name")) == "compatibilityMode":
            compatibility_mode = node
            break
    if compatibility_mode is None:
        compatibility_mode = OxmlElement("w:compatSetting")
        compatibility_mode.set(qn("w:name"), "compatibilityMode")
        compatibility_mode.set(
            qn("w:uri"),
            "http://schemas.microsoft.com/office/word",
        )
        compat.append(compatibility_mode)
    compatibility_mode.set(qn("w:val"), "15")

    # Do not request automatic field updates when Word opens the document.
    # Word Desktop can present an external-field warning even when the only
    # fields are the local PAGE and NUMPAGES footer fields. Word updates those
    # pagination fields during normal pagination/print layout without this flag.
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)


def _set_table_direction(table, rtl: bool) -> None:
    properties = table._tbl.tblPr
    node = properties.find(qn("w:bidiVisual"))
    if node is None:
        node = OxmlElement("w:bidiVisual")
        properties.append(node)
    node.set(qn("w:val"), "1" if rtl else "0")


def _repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    node = properties.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        properties.append(node)
    node.set(qn("w:val"), "1")


def _cant_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def _set_table_geometry(document: Document, table, block: dict, headers: list[str], rows: list[list[str]]) -> None:
    explicit = block.get("width_weights")
    if explicit is not None:
        weights = [float(value) for value in explicit]
        if len(weights) != len(headers) or any(value <= 0 for value in weights):
            raise ValueError("width_weights must contain one positive value per column")
    else:
        columns = list(zip(headers, *rows)) if rows else [(header,) for header in headers]
        weights = [max(1.0, min(6.0, max(len(value) for value in column) / 8.0)) for column in columns]

    section = document.sections[-1]
    total_width = int(round((section.page_width - section.left_margin - section.right_margin) / 635))
    total_weight = sum(weights)
    widths = [int(round(total_width * weight / total_weight)) for weight in weights]
    widths[-1] += total_width - sum(widths)

    table.autofit = False
    properties = table._tbl.tblPr
    for tag, attributes in (
        ("w:tblW", {"w:w": str(total_width), "w:type": "dxa"}),
        ("w:tblInd", {"w:w": "0", "w:type": "dxa"}),
        ("w:tblLayout", {"w:type": "fixed"}),
    ):
        node = properties.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            properties.append(node)
        for key, value in attributes.items():
            node.set(qn(key), value)

    cell_margins = properties.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        properties.append(cell_margins)
    for side, value in (("top", 80), ("start", 100), ("bottom", 80), ("end", 100)):
        node = cell_margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid_columns = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for column, width in zip(grid_columns, widths):
        column.set(qn("w:w"), str(width))
    for row in table.rows:
        _cant_split(row)
        for cell, width in zip(row.cells, widths):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            width_node = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            width_node.set(qn("w:w"), str(width))
            width_node.set(qn("w:type"), "dxa")


def _set_cell_text(cell, text: str, rtl: bool, header: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.style = "Arabic Table Header RTL" if header else "Arabic Table Body RTL"
    _set_paragraph_direction(paragraph, rtl)
    run = paragraph.add_run(str(text))
    _set_run_direction(run, rtl)


def _add_table(document: Document, block: dict) -> None:
    headers = [str(value) for value in block.get("headers", [])]
    rows = [[str(value) for value in row] for row in block.get("rows", [])]
    if not headers:
        raise ValueError("A table requires at least one header")
    if any(len(row) != len(headers) for row in rows):
        raise ValueError("Every table row must match the header column count")

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    rtl = str(block.get("direction", "rtl")) != "ltr"
    _set_table_direction(table, rtl)
    _repeat_header(table.rows[0])
    for cell, value in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, value, rtl, True)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            _set_cell_text(cell, value, rtl)
    _set_table_geometry(document, table, block, headers, rows)


def _add_image(document: Document, block: dict) -> None:
    source = Path(str(block.get("path", "")))
    if not source.is_file():
        raise FileNotFoundError(f"Image not found: {source}")
    section = document.sections[-1]
    text_width = section.page_width - section.left_margin - section.right_margin
    requested = Inches(float(block.get("max_width_inches", 6.5)))
    width = min(text_width, requested)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = paragraph.add_run().add_picture(str(source), width=width)
    alt = str(block.get("alt", "")).strip()
    if alt:
        shape._inline.docPr.set("descr", alt)

    caption = str(block.get("caption", "")).strip()
    if caption:
        caption_paragraph = document.add_paragraph(style="Caption")
        rtl = _is_rtl(caption)
        _set_paragraph_direction(caption_paragraph, rtl)
        _add_directional_runs(caption_paragraph, caption, rtl)


def _configure_section(section, landscape: bool) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    section.page_width = Mm(297 if landscape else 210)
    section.page_height = Mm(210 if landscape else 297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)


def _add_dynamic_field(paragraph, field_name: str) -> None:
    begin_run = paragraph.add_run()
    _set_run_direction(begin_run, False)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    _set_run_direction(instruction_run, False)
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {field_name} "
    instruction_run._r.append(instruction)

    separate_run = paragraph.add_run()
    _set_run_direction(separate_run, False)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run("1")
    _set_run_direction(result_run, False)

    end_run = paragraph.add_run()
    _set_run_direction(end_run, False)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _configure_headers_and_footers(document: Document, title: str) -> None:
    title_rtl = _is_rtl(title)
    for section in document.sections:
        section.header_distance = Mm(8)
        section.footer_distance = Mm(8)
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header = section.header.paragraphs[0]
        header.clear()
        header.style = "Arabic Body RTL" if title_rtl else "English Body LTR"
        _set_paragraph_direction(header, title_rtl)
        _add_directional_runs(header, title, title_rtl)
        header.paragraph_format.space_after = Pt(0)
        for run in header.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(89, 89, 89)

        footer = section.footer.paragraphs[0]
        footer.clear()
        footer.style = "English Body LTR"
        _set_paragraph_direction(footer, False)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.space_after = Pt(0)
        _add_dynamic_field(footer, "PAGE")
        _add_directional_runs(footer, " / ", False)
        _add_dynamic_field(footer, "NUMPAGES")
        for run in footer.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(89, 89, 89)


def build_document(model: dict, output_path: str | Path) -> dict:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_section(document.sections[0], False)
    _configure_styles(document)
    _configure_modern_word_compatibility(document)

    title = str(model.get("title", "")).strip()
    if title:
        _add_text_paragraph(document, title, style_name="Arabic Title RTL" if _is_rtl(title) else "English Body LTR")

    blocks = model.get("blocks", [])
    for index, block in enumerate(blocks):
        block_type = block.get("type")
        if block_type == "paragraph":
            text = str(block.get("text", ""))
            direction = str(block.get("direction", "auto"))
            rtl = _is_rtl(text, direction)
            _add_text_paragraph(
                document,
                text,
                direction,
                str(block.get("style") or ("Arabic Mixed Body RTL" if rtl and re.search(r"[A-Za-z]", text) else "Arabic Body RTL" if rtl else "English Body LTR")),
            )
        elif block_type == "heading":
            text = str(block.get("text", ""))
            level = int(block.get("level", 1))
            if level not in (1, 2):
                raise ValueError("Heading level must be 1 or 2")
            style_name = f"Arabic Heading {level} RTL" if _is_rtl(text) else "English Body LTR"
            _add_text_paragraph(document, text, style_name=style_name)
        elif block_type == "list":
            ordered = bool(block.get("ordered", False))
            for item in block.get("items", []):
                text = str(item)
                rtl = _is_rtl(text)
                style_name = (
                    "Arabic List Number RTL" if rtl and ordered else
                    "Arabic List Bullet RTL" if rtl else
                    "English List Number LTR" if ordered else
                    "English List Bullet LTR"
                )
                _add_text_paragraph(document, text, style_name=style_name)
        elif block_type == "callout":
            kind = str(block.get("kind", "note")).lower()
            style_name = "Warning" if kind == "warning" else "Note"
            _add_text_paragraph(document, str(block.get("text", "")), style_name=style_name)
        elif block_type == "page_break":
            document.add_page_break()
        elif block_type == "table":
            headers = block.get("headers", [])
            landscape = bool(block.get("landscape", len(headers) >= 6))
            if landscape:
                _configure_section(document.add_section(WD_SECTION.NEW_PAGE), True)
            _add_table(document, block)
            if landscape and index < len(blocks) - 1:
                _configure_section(document.add_section(WD_SECTION.NEW_PAGE), False)
        elif block_type == "image":
            _add_image(document, block)

    _configure_headers_and_footers(document, title)
    document.save(output)
    return {"output": str(output), "blocks": len(model.get("blocks", []))}


def main() -> int:
    from doc_model import load_model

    parser = argparse.ArgumentParser(description="Build an Arabic-first or mixed-language DOCX from JSON")
    parser.add_argument("model", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    started = time.perf_counter()
    result = build_document(load_model(args.model), args.output)
    result["elapsed_seconds"] = round(time.perf_counter() - started, 3)
    print(json.dumps(result, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())


__all__ = ["build_document"]
